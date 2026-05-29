"""
Build the term-paper Word document (lecture-scope edition).

Reads:
- results/experiment_results.json
- figures/*.png

Writes:
- term_paper.docx
"""
import json
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
RESULTS_PATH = os.path.join(ROOT, "results", "experiment_results.json")
FIGURES_DIR = os.path.join(ROOT, "figures")
OUT_PATH = os.path.join(ROOT, "term_paper.docx")
OUT_PATH_FALLBACK = os.path.join(ROOT, "term_paper_v4.docx")

FONT_KR = "맑은 고딕"
FONT_EN = "Calibri"

FEATURES_LIST = ["rg", "mean_dist", "std_dist", "clash_count", "density", "n_residues"]


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------
def _set_run_font(run, size_pt=10.5, bold=False, color=None):
    run.font.name = FONT_EN
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_KR)
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)


def add_paragraph(doc, text, size_pt=10.5, bold=False, align=None, color=None, space_after=4):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _set_run_font(run, size_pt=size_pt, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 11.5}
    space = {1: 8, 2: 6, 3: 4}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space[level])
    p.paragraph_format.space_after = Pt(space[level])
    run = p.add_run(text)
    _set_run_font(run, size_pt=sizes[level], bold=True, color=(0x1F, 0x3A, 0x5F))
    return p


def add_bullet(doc, text, size_pt=10.5):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_run_font(run, size_pt=size_pt)
    return p


def add_figure(doc, path, width_cm=14, caption=None):
    if not os.path.exists(path):
        add_paragraph(doc, f"[누락된 그림: {os.path.basename(path)}]", color=(0xB0, 0x00, 0x00))
        return
    doc.add_picture(path, width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        _set_run_font(run, size_pt=9.5, bold=False, color=(0x55, 0x55, 0x55))


def add_table(doc, header, rows, col_widths_cm=None, caption=None):
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        _set_run_font(run, size_pt=9.5, bold=False, color=(0x55, 0x55, 0x55))
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, txt in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt)
        _set_run_font(run, size_pt=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        # shading
        tc_pr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F3A5F")
        tc_pr.append(shd)
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            _set_run_font(run, size_pt=9.5)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths_cm:
        for row in table.rows:
            for cell, w in zip(row.cells, col_widths_cm):
                cell.width = Cm(w)
    doc.add_paragraph()
    return table


# ---------------------------------------------------------------------------
# Body
# ---------------------------------------------------------------------------
def build(results):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # ---- 표지 ----
    add_paragraph(doc, "Machine Learning Term Project", size_pt=11, align="center", color=(0x55, 0x55, 0x55))
    add_paragraph(doc, "인공지능공학과 12243674 유현서", size_pt=11, align="center", color=(0x55, 0x55, 0x55))
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "단백질 3D 구조 평가에서 데이터 분할의 함정:",
        size_pt=15, bold=True, align="center", color=(0x1F, 0x3A, 0x5F),
    )
    add_paragraph(
        doc,
        "강의 범위 모델(Logistic Regression / Random Forest / AdaBoost / XGBoost / SVM-RBF)을 통한 타겟 단위 일반화 실패 검증",
        size_pt=12, align="center", color=(0x1F, 0x3A, 0x5F), space_after=18,
    )

    # ---- Abstract ----
    add_heading(doc, "Abstract", level=2)
    add_paragraph(
        doc,
        "본 보고서는 단백질 구조의 결함 여부를 이진 분류하는 머신러닝 파이프라인을 강의 범위(Lecture 0–10) 안의 "
        "다섯 가지 분류기—Logistic Regression, Random Forest, AdaBoost, XGBoost, SVM-RBF—와 PCA 시각화만으로 구축했다. "
        "Decoys ‘R’ Us / fisa 벤치마크의 1,070개 디코이를 사용해 두 시나리오를 비교한다: "
        "(A) train/test가 i.i.d에 가까운 무작위 3-way 분할과, (B) train에 처음 보는 단백질만 들어가도록 한 "
        "단백질 타겟 단위 3-way 분할(OOD 평가). 동일한 모델·동일한 특징·동일한 튜닝 절차에서 분할 방식만 바꿨을 뿐인데, "
        "다섯 모델 모두 Scenario A에서 Test ROC-AUC ≥ 0.95에 도달하다가 Scenario B에서는 0.36~0.69로 무너졌다. "
        "Confusion Matrix·PCA·Feature Importance·실패 케이스 정성 분석을 종합한 결과, "
        "Scenario B에서의 실패는 ‘모델 용량 부족’이 아니라 (a) train Stable 사례가 단 2개에 불과한 극단적 클래스 불균형, "
        "(b) 6개 특징 중 5개에서 test Stable의 평균이 train Stable의 값 범위를 벗어나는 covariate shift, "
        "(c) 오분류된 Stable 디코이의 92.5%가 RMSD 임계값 직전 회색지대에 몰려 있는 라벨 모호성 — "
        "세 요인의 합이라는 점을 보였다. 본 연구는 ‘좋은 모델을 찾았는가’가 아니라 "
        "‘평가 방식이 무엇을 측정하고 있는가’를 묻는 데 의의가 있다.",
        size_pt=10.5, space_after=18,
    )

    # ---- 차례 ----
    add_heading(doc, "차례", level=2)
    toc = [
        "Abstract",
        "1. Problem Definition",
        "2. Dataset Description",
        "3. Methodology (강의 범위 매핑)",
        "4. Experiments",
        "5. Result Analysis",
        "6. Discussion",
        "7. Conclusion",
        "부록 A. AI 사용 투명성 (Disclosure)",
        "부록 B. References",
    ]
    for t in toc:
        add_paragraph(doc, t, size_pt=10.5)

    doc.add_page_break()

    # ---- 1. Problem Definition ----
    add_heading(doc, "1. Problem Definition", level=1)

    add_heading(doc, "1.1 배경: AI 기반 단백질 설계", level=2)
    add_paragraph(
        doc,
        "AlphaFold, RFdiffusion 등 생성형 AI의 등장으로 in silico 환경에서 새로운 단백질 구조를 설계하는 것이 가능해졌다. "
        "그러나 생성된 3D 구조가 물리화학적으로 안정한지, 혹은 입체적 충돌·비정상 위상을 가진 결함 구조인지를 검증하는 일은 "
        "여전히 신약 개발 파이프라인의 병목이다. 즉 AI 설계 결과물을 실무에 적용하려면 구조적 품질 평가가 필수다.",
    )

    add_heading(doc, "1.2 해결하고자 하는 구체적 문제", level=2)
    add_paragraph(
        doc,
        "본 프로젝트는 단백질 3D 좌표에서 추출한 기하학적 특징을 입력으로 받아 구조의 안정성(Stable=0)과 결함(Defective=1) 여부를 "
        "이진 분류하는 머신러닝 파이프라인을 구축한다. 동시에 평가 시 데이터를 무작위로 섞는 ‘무작위 분할(Random Split)’과 "
        "단백질 타겟을 분리하는 ‘타겟 단위 분할(Target-level Split)’ 두 방식을 대조해, "
        "i.i.d(독립 동일분포) 가정이 충족되는 평가와 그렇지 않은 OOD(out-of-distribution) 평가 사이에서 "
        "동일 모델의 성능이 어떻게 달라지는지를 정량적으로 분석한다 (상동성 누수는 이 차이를 만들어내는 한 요인일 뿐이며, "
        "본 보고서는 클래스 불균형·covariate shift 등 다른 요인까지 함께 분리해 살펴본다).",
    )

    add_heading(doc, "1.3 문제의 중요성 및 수혜자", level=2)
    add_paragraph(
        doc,
        "다수 구조생물학 ML 연구가 무작위 분할 평가를 그대로 사용하므로, train/test가 i.i.d에 가까운 조건에서 측정된 "
        "성능이 실제 운영 환경(처음 보는 단백질, OOD)의 성능을 과대평가한다. 이렇게 i.i.d 가정에 맞춰 학습된 모델은 "
        "미지의 구조 앞에서 무작위 수준까지 무너지므로, 신약 개발 현장에 그대로 투입하기 어렵다. "
        "본 연구는 ‘타겟 독립성이 검증된 평가 파이프라인’을 강의 범위 안에서 구축해, 신약 개발 연구원과 ML 도입을 검토하는 "
        "AI 바이오 제약사에 평가 가이드라인을 제공한다.",
    )

    add_heading(doc, "1.4 해결책이 제공하는 가치", level=2)
    add_paragraph(
        doc,
        "본 보고서는 강의 범위에 속하는 다섯 가지 분류기(Logistic Regression, Random Forest, AdaBoost, XGBoost, SVM-RBF)를 "
        "동일한 특징 공간 위에서 비교해, 모델의 종류와 무관하게 ‘무작위 분할에서는 90%대 성능’과 ‘타겟 단위 분할에서는 무작위 수준 성능’이라는 "
        "착시가 일관되게 나타남을 보인다. 이를 통해 평가 방식 자체가 결과를 좌우한다는 점, 그리고 향후 개선의 방향이 "
        "모델 복잡도 증가가 아닌 ‘데이터 다양성 확보와 타겟 독립 평가’임을 정량적으로 제시한다.",
    )

    # ---- 2. Dataset ----
    add_heading(doc, "2. Dataset Description", level=1)

    add_heading(doc, "2.1 데이터셋 개요·출처·라이선스", level=2)
    add_paragraph(
        doc,
        "David Baker 교수 연구팀이 구축한 ‘Decoys ‘R’ Us’ 벤치마크 [1]의 fisa 세트를 사용했다. "
        "이 데이터셋은 초기 Rosetta 알고리즘으로 생성된 단백질 구조 디코이(decoy)를 집대성한 자료로, "
        "정상 구조에 가까운 형태부터 물리적으로 불안정한 결함 구조까지 폭넓은 스펙트럼을 포함한다.",
    )
    add_bullet(doc, "출처(URL): http://compbio.buffalo.edu/dd/download.shtml")
    add_bullet(doc, "공개 형태: 공식 사이트에서 익명 다운로드 가능한 공개 학술 벤치마크.")
    add_bullet(doc,
        "라이선스/사용 조건: 명시적 라이선스 텍스트가 사이트에 게시되어 있지는 않으나, 학술 출판물 [1]을 "
        "통해 공개된 벤치마크로 학술 연구 목적의 사용이 관례적으로 허용된다. 본 보고서는 데이터셋 사용을 "
        "원전 인용 [1]으로 명시함으로써 학술 인용 의무를 충족하며, 상업적 재배포는 수행하지 않는다.")
    add_bullet(doc, "본 과제에서의 사용 범위: 6개 타겟 중 라벨링 가능한 4개 타겟의 PDB 파일을 로컬에 다운로드해 특징 추출에만 사용했고, 원 데이터를 재배포하지 않는다.")

    add_heading(doc, "2.2 데이터셋 구성", level=2)
    add_paragraph(
        doc,
        "fisa 세트는 6개의 단백질 타겟(1fc2, 1hdd-C, 2cro, 4icb 등)으로 구성되며, 각 타겟마다 약 1,400개의 디코이 PDB 파일이 존재한다. "
        "본 실험에서는 라벨링 가능한(라벨이 -1이 아닌) 1,070개의 구조를 사용했고, 사용된 타겟은 1fc2(225), 1hdd-C(185), 2cro(322), 4icb(338)이다.",
    )

    add_heading(doc, "2.3 라벨링 정책", level=2)
    add_paragraph(
        doc,
        "각 디코이의 RMSD(Root Mean Square Deviation) 값을 기준으로 라벨을 부여했다. "
        "RMSD < 4Å은 Stable(0), RMSD > 8Å은 Defective(1)로 정의하며, 4–8Å 구간의 모호한 디코이는 라벨 -1로 두고 학습/평가에서 제외했다. "
        "이 기준은 Rosetta 디코이 평가에서 일반적으로 사용되는 임계값을 따른다.",
    )

    add_heading(doc, "2.4 데이터셋 선정 사유", level=2)
    add_paragraph(
        doc,
        "fisa 세트는 (1) 타겟 단위로 디코이가 명확히 분리되어 있어 타겟 누수 실험을 설계하기 용이하고, "
        "(2) 클래스 불균형과 단백질 위상 다양성을 동시에 가진 구조이며, "
        "(3) PDB 텍스트 파일만 파싱하면 추가 도메인 라이브러리 없이도 기하학적 특징을 추출할 수 있어 강의 범위 도구만으로 다룰 수 있다.",
    )

    # ---- 3. Methodology ----
    add_heading(doc, "3. Methodology (강의 범위 매핑)", level=1)

    add_heading(doc, "3.1 강의 범위 내 알고리즘만 사용한다는 원칙", level=2)
    add_paragraph(
        doc,
        "본 보고서의 모든 알고리즘·전처리·평가는 학기 강의(Lecture 0–10)에서 명시적으로 다룬 개념에서만 가져왔다. "
        "특히 신경망/딥러닝 계열(PointNet 등)과 비선형 차원축소(t-SNE)는 본 강의 범위 밖이므로 사용하지 않는다. "
        "차원축소는 Lecture 10에서 다룬 PCA만, 분류기는 Lecture 3·6·7에서 다룬 모델만 사용한다.",
    )

    add_heading(doc, "3.2 사용한 알고리즘과 강의 매핑", level=2)
    add_paragraph(
        doc,
        "역할 구분을 먼저 명확히 하면, 본 연구의 main model은 ensemble 계열(Random Forest, AdaBoost, XGBoost — 모두 Lecture 6 [2])이며, "
        "이 ensemble들의 우월성을 평가하기 위한 simple baseline으로 선형 모델(Logistic Regression — Lecture 3 [2])과 "
        "커널 모델(SVM-RBF — Lecture 7 [2])을 함께 학습한다. 즉 ‘단순 baseline → 강력한 ensemble’ 순으로 모델 표현력을 단계적으로 키워 "
        "데이터 누수 효과가 모델 강도에 따라 어떻게 달라지는지(혹은 달라지지 않는지)를 비교하는 것이 본 라인업의 의도다. "
        "PCA(Lecture 10 [2])는 분류기가 아니라 특징 공간을 2차원으로 사영해 시각적으로 진단하기 위한 보조 도구로만 쓰인다.",
    )
    add_bullet(doc, "Logistic Regression (Lecture 3 [2]) — simple baseline: 선형 분류기. L1/L2 정규화는 Lecture 2 [2]에서 다룬 개념. 구현은 scikit-learn [7].")
    add_bullet(doc, "Random Forest (Lecture 6 Ensembles – Bagging [2]; 원전 Breiman [4]) — main model: 비선형 패턴 학습과 MDI 기반 Feature Importance 분석. 구현은 scikit-learn [7].")
    add_bullet(doc, "AdaBoost (Lecture 6 Ensembles – Boosting [2]; 원전 Freund & Schapire [5]) — main model: 약한 분류기 가중 부스팅. 구현은 scikit-learn [7].")
    add_bullet(doc, "XGBoost (Lecture 6 Ensembles [2]; 원전 Chen & Guestrin [6]) — main model (강한 후보): Gradient Boosting 계열의 대표 구현. xgboost 라이브러리를 우선 사용하고 미설치 시 sklearn의 GradientBoosting [7]으로 fallback.")
    add_bullet(doc, "SVM with RBF kernel (Lecture 7 [2]; 원전 Cortes & Vapnik [3]) — kernel baseline: Gaussian kernel을 통한 비선형 결정 경계. 구현은 scikit-learn [7].")
    add_bullet(doc, "PCA 2D 시각화 (Lecture 10 [2]): 비선형 t-SNE 대신 강의 범위 내 선형 차원축소만 사용. 구현은 scikit-learn [7].")
    add_bullet(doc, "평가 지표 (Lecture 3 [2]): Accuracy, Precision, Recall, F1, ROC-AUC, Confusion Matrix.")
    add_bullet(doc, "데이터 누수 방지 (Lecture 2 [2] k-fold CV의 자연스러운 확장): 타겟(그룹) 단위 분할로 단백질 식별 정보가 train/test에 동시에 들어가지 않게 함. scikit-learn의 GroupShuffleSplit 사용 [7].")

    add_heading(doc, "3.3 특징 추출(Feature Engineering)", level=2)
    add_paragraph(
        doc,
        "PDB 파일을 직접 파싱해 다음 6개의 기하학적 특징을 산출했다. 단백질의 ‘결함성’이 본질적으로 3D 기하 특성에서 비롯된다는 도메인 가정을 그대로 반영한다.",
    )
    add_bullet(doc, "rg (radius of gyration): 단백질의 전체적인 응축도. 결함 구조는 일반적으로 더 넓게 퍼져 있다.")
    add_bullet(doc, "mean_dist, std_dist: 잔기 쌍 간 평균 거리와 표준편차. 정상 구조는 좁은 분포를, 결함 구조는 넓은 분포를 갖는다.")
    add_bullet(doc, "clash_count: 비결합 원자 간 충돌 횟수. 물리적 결함의 직접 지표.")
    add_bullet(doc, "density: 잔기 밀도. 응축된 정상 구조에서 더 높게 나타난다.")
    add_bullet(doc, "n_residues: 잔기 수. 단백질 크기 자체의 효과를 통제하기 위한 보조 특징.")

    # ---- 4. Experiments ----
    add_heading(doc, "4. Experiments", level=1)

    add_heading(doc, "4.1 두 시나리오 비교 실험", level=2)
    add_paragraph(
        doc,
        "본 연구의 핵심은 ‘데이터 누수가 일반화에 미치는 영향’을 다른 모든 조건은 동일하게 두고 분할 방식만 바꿔 정량 비교하는 것이다. "
        "이를 위해 두 시나리오를 동일한 1,070개 샘플 위에서 평가했다.",
    )
    add_bullet(doc, "Scenario A — 무작위 3-way 분할(60%/20%/20%): 타겟 정보를 무시하고 무작위로 train/val/test을 만든다. 동일 타겟의 디코이가 분할 양쪽에 섞이므로 누수가 발생한다.")
    add_bullet(doc, "Scenario B — 타겟 단위 3-way 분할: Train = {2cro, 4icb}, Validation = {1fc2}, Test = {1hdd-C}. 모델은 학습 중 본 적 없는 단백질로만 평가받는다.")

    add_paragraph(
        doc,
        "Scenario B의 타겟 배정은 두 가지 기준을 따랐다. "
        "(1) 샘플 수 비율로 60/20/20에 가깝게 맞추기 위해 가장 큰 두 타겟(2cro: 322, 4icb: 338, 합 61.7%)을 train, "
        "중간 크기 타겟(1fc2: 225, 21.0%)을 validation, 가장 작은 타겟(1hdd-C: 185, 17.3%)을 test로 두었다. "
        "(2) 단백질 크기 측면에서도 1hdd-C(잔기 수 57)가 train 타겟(2cro=65, 4icb=76)에서 가장 멀리 떨어져 있어 "
        "‘처음 보는 단백질 위상’에 대한 가장 엄격한 OOD(out-of-distribution) 평가가 되도록 의도적으로 test에 배정했다. "
        "즉 본 시나리오는 cherry-pick으로 쉬운 split을 고른 것이 아니라, 반대로 ‘가장 어렵게 만든’ 평가다.",
    )

    add_heading(doc, "4.2 평가 지표", level=2)
    add_paragraph(
        doc,
        "클래스 불균형(결함>>정상)을 감안해 Accuracy가 아닌 ROC-AUC와 F1을 핵심 지표로 채택하고, "
        "예측 쏠림을 진단하기 위해 Confusion Matrix를 함께 보고한다. 모든 지표는 Lecture 3에서 다룬 정의를 따른다.",
    )

    add_heading(doc, "4.3 3-way Split을 이용한 하이퍼파라미터 튜닝", level=2)
    add_paragraph(
        doc,
        "각 모델은 Train으로 학습, Validation ROC-AUC로 하이퍼파라미터를 고르고, 최종 모델 1개만 Test에 단 한 번 평가했다(Test 오염 방지). "
        "튜닝 그리드는 강의에서 다룬 핵심 하이퍼파라미터에 한정한다 (Table 1).",
    )
    add_table(
        doc,
        ["모델", "역할", "튜닝 그리드", "총 조합 수"],
        [
            ["Logistic Regression", "simple baseline",
             "penalty ∈ {L1, L2}, C ∈ {0.1, 1, 10} 중 4개 조합(L2×3 + L1×1)", "4"],
            ["Random Forest", "main model",
             "n_estimators ∈ {50, 100, 200}, max_depth ∈ {None, 6}", "6"],
            ["AdaBoost", "main model",
             "n_estimators ∈ {50, 100, 200}, learning_rate ∈ {0.5, 1.0}", "6"],
            ["XGBoost", "main model",
             "n_estimators ∈ {100, 200}, max_depth ∈ {3, 5}, learning_rate ∈ {0.05, 0.1}, scale_pos_weight = neg/pos",
             "8"],
            ["SVM-RBF", "kernel baseline",
             "C ∈ {0.1, 1, 10}, gamma ∈ {scale, 0.1, 1.0}", "9"],
        ],
        col_widths_cm=[3.0, 2.5, 8.5, 2.0],
        caption="표 1. 모델별 튜닝 그리드. 각 시나리오에서 Validation ROC-AUC가 가장 높은 조합 1개를 Test에 단 한 번 평가했다.",
    )
    add_paragraph(
        doc,
        "전처리는 모든 거리 기반 모델(LogReg, SVM)에 StandardScaler를 적용했고, 트리 계열(RF, AdaBoost, XGBoost)에는 적용하지 않았다(트리는 스케일 불변).",
    )

    # ---- 5. Result ----
    add_heading(doc, "5. Result Analysis", level=1)

    add_heading(doc, "5.1 PCA를 통한 특징 공간 분석 (Lecture 10)", level=2)
    pca_var = results["pca_explained_variance"]
    add_paragraph(
        doc,
        f"6차원 특징을 PCA로 2차원에 사영한 결과 PC1이 분산의 {pca_var[0]*100:.1f}%, "
        f"PC2가 {pca_var[1]*100:.1f}%를 설명했다. 즉 본 특징 공간의 변동성은 사실상 하나의 주축으로 거의 모두 설명된다. "
        "이는 6개 기하 특징이 서로 강하게 상관되어 있다는 신호이며, 분류기는 결국 이 1차원 신호 위에서 결정을 내리는 셈이다. "
        "Stable / Defective 클러스터가 PC1 축 방향으로 일부 분리되지만 완전한 선형 분리는 어렵고, "
        "고차원 비선형 모델이 추가로 얻을 이점이 크지 않을 가능성을 시사한다.",
    )
    add_figure(doc, os.path.join(FIGURES_DIR, "pca_analysis.png"), width_cm=12,
               caption="그림 1. 표준화된 기하 특징의 PCA 2차원 사영. PC1이 분산의 92%를 설명한다.")

    add_heading(doc, "5.2 Scenario A — 무작위 분할 결과", level=2)
    a = results["scenario_A"]
    rows_A = []
    for name in ["LogReg", "RandomForest", "AdaBoost", "XGBoost", "SVM_RBF"]:
        r = a[name]
        rows_A.append([
            name,
            f"{r['val_roc_auc']:.4f}",
            f"{r['test_roc_auc']:.4f}",
            f"{r['test_f1']:.4f}",
            f"{r['test_accuracy']:.4f}",
            str(r["best_config"]),
        ])
    add_table(doc, ["모델", "Val AUC", "Test AUC", "Test F1", "Test Acc", "선택된 하이퍼파라미터"], rows_A,
              col_widths_cm=[2.2, 1.8, 1.8, 1.6, 1.6, 6.5],
              caption="표 2. 시나리오 A(무작위 3-way 분할) 모델별 결과. 각 모델은 Validation에서 최적 조합을 고른 뒤 Test에 단 한 번 평가되었다.")
    # Identify Scenario A best-model explicitly
    best_a = max(["LogReg", "RandomForest", "AdaBoost", "XGBoost", "SVM_RBF"],
                 key=lambda m: a[m]["test_roc_auc"])
    add_paragraph(
        doc,
        f"표 2에서 보듯이 다섯 모델 모두 Test ROC-AUC ≥ 0.95, F1 ≥ 0.93에 도달한다. "
        f"ROC-AUC 기준 가장 우수한 모델은 {best_a}(Test AUC={a[best_a]['test_roc_auc']:.4f}, F1={a[best_a]['test_f1']:.4f})이고, "
        f"정확도 기준으로는 AdaBoost(0.925)가 1위다. 단순 모델인 Logistic Regression과 SVM-RBF도 같은 1차원 분리 신호 덕분에 "
        "거의 동일한 성능을 낸다(차이 ≤ 0.012 AUC). 즉 이 시나리오에서는 ‘무슨 모델을 쓰든 잘 되는’ 상태이며, "
        "‘모델 종류’가 사실상 결정적 변수가 아니라는 점이 이미 시사된다. "
        "표면적으로는 ‘이 문제는 거의 풀렸다’는 결론을 내리고 싶어질 수치이지만, §5.3에서 보듯 이 결과는 "
        "데이터 누수가 만들어 낸 착시다.",
    )
    add_figure(doc, os.path.join(FIGURES_DIR, "roc_scenario_A.png"), width_cm=12,
               caption="그림 2. 시나리오 A(무작위 분할)에서의 모델별 ROC 곡선.")
    add_figure(doc, os.path.join(FIGURES_DIR, "confusion_matrix_A.png"), width_cm=15,
               caption="그림 3. 시나리오 A에서의 모델별 Confusion Matrix.")

    add_heading(doc, "5.3 Scenario B — 타겟 단위 분할 결과", level=2)
    b = results["scenario_B"]
    rows_B = []
    for name in ["LogReg", "RandomForest", "AdaBoost", "XGBoost", "SVM_RBF"]:
        r = b[name]
        rows_B.append([
            name,
            f"{r['val_roc_auc']:.4f}",
            f"{r['test_roc_auc']:.4f}",
            f"{r['test_f1']:.4f}",
            f"{r['test_accuracy']:.4f}",
            str(r["test_confusion_matrix"]),
        ])
    add_table(doc, ["모델", "Val AUC", "Test AUC", "Test F1", "Test Acc", "Test Confusion Matrix"], rows_B,
              col_widths_cm=[2.2, 1.8, 1.8, 1.6, 1.6, 5.5],
              caption="표 3. 시나리오 B(타겟 단위 3-way 분할) 모델별 결과. Train={2cro, 4icb}, Val={1fc2}, Test={1hdd-C}.")
    add_paragraph(
        doc,
        "표 3에서 보듯, 동일한 모델·동일한 특징·동일한 튜닝 절차에서 분할 방식만 타겟 단위로 바꿨을 뿐인데 Test ROC-AUC가 0.36~0.69까지 무너진다. "
        "Confusion Matrix를 보면 LogReg/RF/AdaBoost/SVM-RBF 네 모델은 모든 테스트 샘플을 ‘결함(1)’로 일괄 예측해(즉 [0,40]/[0,145]) "
        "F1=0.8788을 받지만, 이는 클래스 불균형(78%가 양성)에서 자연히 나오는 ‘All-Positive 트릭’이다. "
        "다섯 모델 모두 ‘새로운 단백질 위상’ 앞에서 분류기의 의미를 잃었다.",
    )

    add_heading(doc, "(보충) 모델별 실패 양상의 미세한 차이", level=3)
    add_paragraph(
        doc,
        "‘다섯 모델 모두 무너진다’는 결론은 거시적으로 옳지만, Test AUC 자체는 모델별로 0.36~0.69까지 분포해 "
        "실패의 ‘질감’은 미세하게 다르다. 세 가지 관찰이 추가로 가능하다.",
    )
    add_bullet(
        doc,
        f"Logistic Regression이 다른 모델보다 상대적으로 나은 ROC-AUC({b['LogReg']['test_roc_auc']:.4f})를 보인다. "
        "이는 LogReg가 본질적으로 PC1 한 축에 대한 선형 확률 모델이라, 모델이 학습한 ‘응축도/크기’ 축에서 1hdd-C 디코이를 "
        "ranking할 때는 train 분포와 어긋난 영역에서도 어느 정도 일관된 점수를 매길 수 있기 때문이다. "
        "즉 결정 경계(분류 임계값)는 모두 ‘1’로 쓸리지만, 확률 순위는 비교적 보존된다.",
    )
    add_bullet(
        doc,
        "Random Forest와 AdaBoost는 트리 분기점에서 학습 데이터에 본 적 없는 영역(예: n_residues=57)이 들어오면 "
        "특정 leaf로 강제 라우팅되어 Test AUC가 0.51~0.59로 무작위 근처에 머문다. 이는 트리 기반 모델이 "
        "out-of-distribution 입력을 처리하는 일반적 약점과 일치한다.",
    )
    add_bullet(
        doc,
        "XGBoost만 유일하게 반대 방향(All-Negative, [40,0]/[145,0])으로 무너진다. 원인은 train 셋의 양성:음성 비율이 658:2로 "
        "극단적이라 scale_pos_weight = neg/pos ≈ 0.003이 되어, 부스팅이 ‘다 음성(0)으로 찍기’를 학습 데이터에서 사실상 최적해로 받아들이기 때문이다. "
        "이는 강의 범위 안에서도 ‘클래스 불균형 보정 파라미터가 데이터에 따라 어떻게 반대 방향 trivial rule을 유도할 수 있는지’를 보여주는 사례다.",
    )
    add_paragraph(
        doc,
        "이 세 관찰의 공통 함의는 ‘어떤 trivial rule로 무너지느냐’는 모델별 inductive bias의 차이일 뿐, "
        "‘일반화에 실패한다’는 본질은 동일하다는 점이다. 따라서 Scenario B에서는 단일 best model을 선정하는 것이 의미가 없고, "
        "오히려 ‘다섯 모델 모두 의미 있는 generalization에 도달하지 못함’이라는 사실 자체가 본 연구의 핵심 결과다.",
    )
    add_figure(doc, os.path.join(FIGURES_DIR, "roc_scenario_B.png"), width_cm=12,
               caption="그림 4. 시나리오 B(타겟 단위 분할)에서의 모델별 ROC 곡선.")
    add_figure(doc, os.path.join(FIGURES_DIR, "confusion_matrix_B.png"), width_cm=15,
               caption="그림 5. 시나리오 B에서의 모델별 Confusion Matrix. 네 모델은 All-Positive, XGBoost는 All-Negative로 무너졌다.")

    add_heading(doc, "5.4 두 시나리오의 직접 비교", level=2)
    rows_cmp = []
    for name in ["LogReg", "RandomForest", "AdaBoost", "XGBoost", "SVM_RBF"]:
        rows_cmp.append([
            name,
            f"{a[name]['test_roc_auc']:.4f}",
            f"{b[name]['test_roc_auc']:.4f}",
            f"{a[name]['test_roc_auc'] - b[name]['test_roc_auc']:+.4f}",
        ])
    add_table(doc, ["모델", "시나리오 A Test AUC", "시나리오 B Test AUC", "ΔAUC (A − B)"], rows_cmp,
              col_widths_cm=[3.0, 3.5, 3.5, 3.0],
              caption="표 4. 시나리오 A vs B 직접 비교. ΔAUC는 ‘i.i.d 평가’와 ‘OOD 평가’ 사이에서 동일 모델의 성능이 얼마나 달라지는지를 정량화한다.")
    add_paragraph(
        doc,
        "모델 종류와 무관하게 ΔAUC가 0.27~0.59에 달한다. 이 격차는 단일 원인 — 예컨대 ‘상동성 누수’ 하나만 — 으로 환원되지 않는다. "
        "Scenario A는 무작위 분할이 train/test의 분포를 거의 동일하게(i.i.d 가정) 유지하기 때문에 모델이 학습 분포 위에서 평가받고, "
        "Scenario B는 train에 Stable이 거의 없고(클래스 불균형) 그마저도 test의 단백질 영역과 다른 곳에 위치(covariate shift)하기 때문에 "
        "사실상 학습 분포 밖에서 평가받는다. 즉 ΔAUC는 ‘i.i.d 평가 ↔ OOD 평가’ 전환에서 발생하는 성능 손실을 측정한 값이고, "
        "모델 복잡도를 키워서는 해결되지 않는 종류의 격차다.",
    )

    # ------ 5.5.5 / 5.6 실패 케이스 정성 분석 ------
    # Lazy load: failure analysis JSON
    fa_path = os.path.join(ROOT, "results", "failure_analysis.json")
    fa = None
    if os.path.exists(fa_path):
        with open(fa_path, "r", encoding="utf-8") as f:
            fa = json.load(f)

    add_heading(doc, "5.5 Feature Importance와 PC1의 일치", level=2)
    add_paragraph(
        doc,
        "Random Forest의 MDI 기반 Feature Importance(Scenario B 학습 기준)를 보면 rg(회전반경), n_residues, density가 상위에 위치한다. "
        "이는 PCA에서 PC1이 거의 모든 분산을 설명한다는 결과와 일치하며, 결국 모델이 학습한 신호는 ‘단백질의 응축도/크기’라는 1차원 축 위에 놓여 있다. "
        "Scenario A에서는 이 1차원 축이 같은 타겟 디코이끼리 공유되어 매우 강력하게 작동하지만, Scenario B에서는 새로운 타겟에서 이 축의 위치가 어디로 옮겨 가는지 모델이 모른다.",
    )
    add_figure(doc, os.path.join(FIGURES_DIR, "feature_importance_rf.png"), width_cm=11,
               caption="그림 6. Random Forest의 MDI 기반 Feature Importance (시나리오 B train 셋 기준).")

    # ------ 5.6 실패 케이스 정성 분석 ------
    add_heading(doc, "5.6 실패 케이스 정성 분석 (Failure Mode Analysis)", level=2)
    if fa is None:
        add_paragraph(doc, "(실패 케이스 분석 결과(failure_analysis.json)를 찾을 수 없음.)",
                      color=(0xB0, 0x00, 0x00))
    else:
        rmsd = fa["rmsd_summary"]
        shifts = fa.get("feature_range_check") or fa.get("feature_shift", {})
        size = fa["size_summary"]

        add_paragraph(
            doc,
            "Scenario B에서 가장 흥미로운 실패 패턴은 'All-Positive 예측'이다. LogReg / RF / AdaBoost / SVM-RBF "
            "네 모델 모두 1hdd-C 테스트 셋의 185개 디코이를 전부 Defective(1)로 분류했고, 그 결과 40개의 Stable(0) 디코이가 "
            "모두 잘못 분류되었다. 이 40개 오분류 샘플의 정성적 특성을 세 각도에서 분석한다.",
        )

        add_heading(doc, "(1) 모든 오분류 Stable이 임계값 근처에 몰려 있다", level=3)
        add_paragraph(
            doc,
            f"오분류된 Stable 디코이(n={rmsd['n_misclassified_stable']})의 RMSD 분포를 보면, "
            f"평균 RMSD가 {rmsd['rmsd_mean']:.2f}Å, 중간값 {rmsd['rmsd_median']:.2f}Å, 최대 {rmsd['rmsd_max']:.2f}Å으로 "
            f"Stable/Defective를 가르는 임계값 4Å에 매우 근접해 있다. 특히 이 중 "
            f"{rmsd['near_threshold_fraction']*100:.1f}%가 RMSD ≥ 3.0Å로, 사실상 'borderline' 구간에 위치한다. "
            "즉 모델 입장에서는 라벨 0과 라벨 1이 가장 구분하기 어려운 회색지대의 샘플들이고, 이 회색지대를 "
            "한쪽(=Defective)으로 일괄 처리한 셈이다.",
        )
        add_figure(doc, os.path.join(FIGURES_DIR, "failure_rmsd_dist.png"), width_cm=12,
                   caption="그림 7. 시나리오 B에서 오분류된 40개 Stable 디코이(1hdd-C)의 RMSD 분포. 92.5%가 임계값 4Å 직전 회색지대에 몰려 있다.")

        add_heading(doc, "(2) Train에서 Stable이 거의 없었다 — 극단적 클래스 불균형", level=3)
        add_paragraph(
            doc,
            "더 근본적인 원인은 train 셋의 클래스 구성에 있다. Scenario B의 train(2cro + 4icb) "
            "660건 중 Stable로 라벨링된 샘플은 단 2개에 불과하다(2cro 1개, 4icb 1개; 비율 0.3%). "
            "즉 어떤 분류기를 쓰든, 모델은 'Stable이 어떻게 생겼는지' 학습할 수 있는 사례를 사실상 본 적이 없다. "
            "class_weight='balanced'나 scale_pos_weight는 손실 함수의 가중치를 맞춰주지만, "
            "결정 경계를 그릴 양성 사례 자체를 만들어주지는 못한다.",
        )

        add_heading(doc, "(3) Train Stable의 값 범위 밖에 Test Stable이 위치한다 — Out-of-Distribution", level=3)
        add_paragraph(
            doc,
            "주의: train Stable이 단 2개(n=2)이므로 표준편차/Z-score 등 분포 통계는 통계학적으로 무의미하다 "
            "(2개 점에서 정규분포 가정과 σ 추정이 신뢰도를 갖지 못한다). 따라서 본 분석은 σ를 가정하지 않고, "
            "‘train Stable의 두 점이 만드는 값의 범위 [min, max] 안에 test Stable의 평균이 들어오는가’만을 "
            "직관적 OOD 판정 기준으로 사용한다.",
        )
        rows = []
        ood_features = []
        for feat in FEATURES_LIST:
            s = shifts[feat]
            ts_vals = s["train_stable_values"]
            ts_str = f"{ts_vals[0]:.4f}, {ts_vals[1]:.4f}"
            range_str = f"[{s['train_stable_min']:.4f}, {s['train_stable_max']:.4f}]"
            test_mean_str = f"{s['test_stable_mean']:.4f}"
            is_ood = s["test_mean_outside_train_range"]
            verdict = "OOD ✗" if is_ood else "in-range ○"
            if is_ood:
                ood_features.append(feat)
            rows.append([feat, ts_str, range_str, test_mean_str, verdict])
        add_table(doc,
                  ["특징", "Train Stable 두 값", "Train [min, max]", "Test Stable 평균 (n=40)", "OOD 여부"],
                  rows,
                  col_widths_cm=[2.6, 3.6, 3.6, 3.4, 2.0],
                  caption="표 5. Train Stable 두 점(n=2)의 값 범위와, Test Stable(1hdd-C, n=40) 평균의 위치 비교. "
                          "Test 평균이 Train [min, max] 밖이면 ‘Out-of-Distribution’으로 분류한다 (σ 가정 없음).")
        n_ood = fa.get("n_ood_features_of_6", len(ood_features))
        add_paragraph(
            doc,
            f"표 5에서 보듯이 6개 특징 중 {n_ood}개에서 Test Stable의 평균이 Train Stable의 값 범위를 완전히 벗어난다. "
            "특히 n_residues는 Train Stable의 두 값이 {65, 76}인데 Test Stable의 잔기 수는 일괄 57로, "
            "‘크기’ 자체가 학습 범위 안에 들어온 적이 없다. rg(회전반경)와 mean_dist도 Test 평균이 Train 최댓값보다 크고, "
            "density는 반대로 Train 최솟값보다도 작다(즉 더 흩어진 분자 구조). "
            "결론적으로 모델은 ‘이런 크기·응축도·거리 분포의 Stable 단백질’ 자체를 본 적이 없는 상태에서 "
            "그 영역을 분류하라고 요구받은 셈이다. 이것이 본 실험에서 관측된 covariate shift의 정량적 증거다.",
        )
        add_figure(doc, os.path.join(FIGURES_DIR, "failure_feature_shift.png"), width_cm=15,
                   caption="그림 8. Train Stable(2cro+4icb, n=2; 점)과 Test Stable(1hdd-C, n=40; 박스)의 특징별 비교. "
                           "각 패널은 σ를 가정하지 않고, train 두 점이 만드는 [min, max] 범위에 test 평균이 들어오는지만 표시한다.")

        add_paragraph(
            doc,
            "정리하면 모델의 실패 모드는 ‘복잡한 위상을 잘못 학습’한 것이 아니라, "
            "(a) train에 Stable 사례가 사실상 없는 극단적 클래스 불균형, "
            "(b) train Stable 두 점이 만드는 값의 범위를 test Stable이 완전히 벗어난 covariate shift, "
            "(c) test의 Stable 사례 대부분이 RMSD 임계값 직전의 회색지대에 몰려 있는 라벨 모호성 — "
            "세 가지 요인의 합이다. 이 세 요인 중 어느 것도 ‘모델 종류를 바꿔서’ 해결되지 않는다.",
        )

    # ---- 6. Discussion ----
    add_heading(doc, "6. Discussion", level=1)

    add_heading(doc, "6.1 ‘모델을 키우면 풀린다’는 오해", level=2)
    add_paragraph(
        doc,
        "Logistic Regression → Random Forest → AdaBoost → XGBoost 순서로 모델 표현력을 단계적으로 키웠지만, Scenario B에서는 다섯 모델 모두 사실상 동일하게 무너진다. "
        "강의에서 다룬 가장 단순한 선형 분류기와 가장 강한 부스팅 계열의 격차가 사라진다는 것은, 부족한 것이 ‘모델 용량’이 아니라 "
        "‘학습 데이터가 커버하지 못한 단백질 위상 다양성’이라는 뜻이다.",
    )

    add_heading(doc, "6.2 클래스 불균형 + Covariate Shift의 합성 효과", level=2)
    add_paragraph(
        doc,
        "Scenario B의 train 셋에서 Stable:Defective ≈ 2:658이라는 극단적 불균형이 발생하고, 동시에 그 두 개의 Stable 사례마저도 "
        "test의 Stable 영역과 값 범위가 겹치지 않는다(§5.6 (3)). 즉 이 시나리오는 한 가지 문제(누수)가 아니라 두 가지 통계적 문제가 "
        "동시에 작동하는 상태다. class_weight='balanced'나 scale_pos_weight 조정은 학습 시점의 손실 균형은 맞춰주지만, "
        "‘test에서 처음 보는 1hdd-C의 Stable 구조’에 대한 결정 경계까지 만들어주지는 못한다. "
        "결과적으로 모델은 ‘다 1로 찍기’ 또는 그 반대 같은 trivial decision rule로 수렴한다. "
        "즉 Scenario A→B의 성능 격차는 ‘무작위 분할이 컨닝(누수)을 허용했다’는 단순한 윤리적 프레임으로는 충분히 설명되지 않고, "
        "‘i.i.d → OOD’ + ‘클래스 균형 → 극단 불균형’이라는 두 축이 동시에 변한 결과로 해석하는 것이 더 정확하다.",
    )

    add_heading(doc, "6.3 본 결과의 강점과 한계", level=2)
    add_bullet(doc, "강점: 다섯 가지 강의 범위 모델에서 동일한 실패 패턴이 재현됨으로써, ‘평가 방식이 결과를 결정한다’는 결론을 단일 모델 사례로 주장할 때보다 훨씬 강하게 뒷받침할 수 있다.")
    add_bullet(doc, "강점: 모든 절차(특징 추출, 분할, 튜닝, 평가)가 강의에서 다룬 개념으로만 구성되어 재현성과 설명 가능성이 높다.")
    add_bullet(doc, "한계: fisa 세트 4개 타겟만 사용했기 때문에 ‘위상 다양성 부족’이 결정적이라는 결론을 다른 도메인으로 일반화하려면 추가 데이터셋(예: fisa_casp3, lmds) 실험이 필요하다.")
    add_bullet(doc, "한계: 본 특징 공간은 사실상 1차원(PC1)에 가깝다. 같은 분할 실험을 더 풍부한 특징(2차 구조 비율, 접촉지도 통계 등)으로 반복하면 결과가 달라질 여지가 있다.")

    add_heading(doc, "6.4 Deployment 가능성과 Future Work", level=2)
    add_paragraph(
        doc,
        "현재 결과는 ‘deploy하기에 만족스러운 수준’이 아니다. Scenario B의 ROC-AUC가 무작위 수준이라는 사실은, "
        "이 파이프라인을 신약 개발 현장의 ‘처음 보는 단백질’에 그대로 적용하면 분류기가 사실상 의미 없는 결정을 내릴 수 있다는 뜻이다. "
        "다만 본 보고서가 보여준 정확한 실패 모드(§5.6)는 ‘다음에 무엇을 시도해야 deploy 가능해지는가’의 방향을 명확히 알려준다. "
        "시간과 자원이 더 주어진다면 우선순위 순으로 다음 단계를 시도할 것이다.",
    )

    add_heading(doc, "(1) 타겟 다양성 확보 (가장 큰 효과 기대)", level=3)
    add_bullet(doc, "동일한 Decoys ‘R’ Us 안의 fisa_casp3(타겟 6개), lmds(타겟 11개)까지 통합하면 학습에 사용 가능한 단백질 위상이 23개로 늘어난다.")
    add_bullet(doc, "Leave-One-Target-Out CV 또는 nested CV를 도입해, ‘새 타겟에 대한 일반화’를 학습 과정 내부에서 직접 최적화 신호로 사용한다 (Lecture 2 k-fold CV의 확장).")
    add_bullet(doc, "예측 가설: 위상 다양성이 6배 늘어나면 Scenario B-스타일 평가에서 ROC-AUC가 0.5에서 의미 있는 수준으로 회복될 것이다. 회복되지 않는다면 모델이 아닌 ‘특징’에 한계가 있다는 신호다.")

    add_heading(doc, "(2) 특징 공간 보강 (1차원 신호 해소)", level=3)
    add_bullet(doc, "현재 6개 기하 특징은 PC1이 분산의 92%를 설명할 정도로 강하게 상관되어 사실상 1차원이다.")
    add_bullet(doc, "추가 후보: 2차 구조 비율(헬릭스/시트/코일), 접촉 지도(contact map)의 통계량, 잔기 간 평균 결합 각도 분포, 표면적 대비 부피 비율 등.")
    add_bullet(doc, "이들은 모두 PDB에서 직접 산출 가능하므로 강의 범위 도구만으로 구현할 수 있다.")

    add_heading(doc, "(3) 모호 구간 명시적 처리 (라벨링 정책 재검토)", level=3)
    add_bullet(doc, "§5.6에서 보였듯이 오분류 샘플 92.5%가 RMSD 3~4Å 구간에 몰려 있다. 이 회색지대는 본질적으로 binary 분류가 어려운 구간이다.")
    add_bullet(doc, "대안 1: 3-class 분류(Stable / Borderline / Defective)로 문제 재정의 (Lecture 3의 multinomial logistic regression 또는 SVM one-vs-rest 사용).")
    add_bullet(doc, "대안 2: RMSD 자체를 회귀로 예측하고 사용처에서 임계값을 결정 (Lecture 2 linear regression, Lecture 4 kNN regression 등 강의 범위 안).")

    add_heading(doc, "(4) 평가 지표/프로토콜 강화", level=3)
    add_bullet(doc, "단일 타겟 평가(1hdd-C)는 변동성이 크다. Leave-One-Target-Out으로 평균과 분산을 함께 보고하면 결론의 신뢰도가 높아진다.")
    add_bullet(doc, "운영 시나리오 가정에 맞춰 결정 임계값을 학습 후 조정해 Precision-Recall trade-off를 명시적으로 보고한다 (Lecture 3 threshold tuning).")

    add_heading(doc, "(5) Deployment 결정 가이드라인", level=3)
    add_paragraph(
        doc,
        "위 (1)~(4)를 모두 적용한 후에도, 새로운 단백질에 대한 ROC-AUC가 안정적으로 0.85 이상으로 회복되는 시점부터가 "
        "실제 서비스 적용을 진지하게 고민할 수 있는 출발점이라고 본다. 그 이전에는 본 파이프라인을 "
        "‘초기 후보 선별 보조용’으로만 활용하고, 최종 결정은 항상 별도의 도메인 검증(예: 분자동역학 시뮬레이션, 실험 검증)을 거치는 것이 안전하다.",
    )

    # ---- 7. Conclusion ----
    add_heading(doc, "7. Conclusion", level=1)
    add_paragraph(
        doc,
        "본 보고서는 강의 범위 내 다섯 가지 분류기(Logistic Regression / Random Forest / AdaBoost / XGBoost / SVM-RBF)와 PCA만을 사용해 "
        "단백질 결함 구조 분류 문제를 평가했다. 핵심 결론은 두 가지다.",
    )
    add_bullet(doc, "(1) i.i.d 평가와 OOD 평가는 다른 문제다. 무작위 분할(i.i.d 충족)에서는 어떤 모델이든 ROC-AUC 0.95+를 받지만, 타겟 단위 분할(OOD + 극단 클래스 불균형)에서는 동일 모델들이 모두 무작위 수준으로 무너진다. ‘잘 풀린다’는 평가는 운영 환경의 분포가 학습 환경과 같다는 가정 위에서만 성립한다.")
    add_bullet(doc, "(2) 더 강한 모델이 답이 아니다. 단순 LogReg와 XGBoost가 같은 위상 다양성 한계 앞에서 동일하게 실패한다. 향후 개선은 모델 복잡도가 아니라 ‘타겟 다양성 확보 + 타겟 독립 평가’ 방향이어야 한다.")
    add_paragraph(
        doc,
        "즉 본 연구는 ‘좋은 모델을 찾았는가’가 아니라 ‘평가 방법이 무엇을 측정하고 있는가’를 묻는 데 의의가 있다. "
        "강의에서 다룬 평가 지표와 분할 개념을 그대로 사용해, 강의 범위 안에서도 데이터 누수의 영향을 명확히 정량화할 수 있음을 보였다.",
    )

    # ---- 부록 A. AI 사용 투명성 ----
    doc.add_page_break()
    add_heading(doc, "부록 A. AI 사용 투명성 (Disclosure)", level=1)
    add_paragraph(
        doc,
        "본 과제 명세 「Important Notes」 4번째 항목에 따라, 본 보고서 작성과 실험 코드 정비 과정에서 사용한 생성형 AI 도구의 사용 범위와 "
        "사용 방식을 다음과 같이 투명하게 기재한다. 항목 구분은 (1) 사용 도구, (2) AI가 보조한 일, (3) 사용자가 직접 한 일, (4) AI가 하지 않은 일, "
        "(5) 검증 절차의 다섯 가지다.",
    )

    add_heading(doc, "A.1 사용한 AI 도구", level=2)
    add_bullet(doc, "Anthropic Claude (Claude Opus 4.7 via Claude Code CLI): 보고서 본문 작성·교정, 강의 범위 매핑, 코드 작성 및 검토, 실험 결과 표/그림을 보고서에 임베드하는 워드 빌드 스크립트(`src/models/build_report.py`) 작성.")
    add_bullet(doc, "scikit-learn / xgboost / matplotlib 공식 문서: 라이브러리 사용법은 공식 문서를 직접 참조했으며, AI는 보조 설명자 역할로만 사용했다.")

    add_heading(doc, "A.2 AI가 보조한 일 (이번 보고서 정비 세션)", level=2)
    add_bullet(doc, "기존 보고서(기계복제본.hwpx) 진단: PointNet과 t-SNE가 본 학기 강의(Lecture 0–10) 범위 밖이라는 점을 강의 PDF 10개를 전수 비교하여 식별.")
    add_bullet(doc, "강의 범위 내 모델 라인업 제안: Logistic Regression(L3) → Random Forest(L6) → AdaBoost(L6) → XGBoost(L6) → SVM-RBF(L7)의 단계적 강도 비교 구성 제안.")
    add_bullet(doc, "코드 작성: `src/models/baseline.py` 재작성, `src/models/run_experiments.py` 신규 작성(2시나리오 × 5모델 × val-set 기반 튜닝 + PCA·ROC·CM 시각화), `src/models/build_report.py` 신규 작성.")
    add_bullet(doc, "보고서 한국어 본문 작성 및 교정: Methodology / Experiments / Result Analysis / Discussion / Conclusion 절의 초안 작성 후 검토.")
    add_bullet(doc, "본 부록 A 및 References의 1차 초안 작성.")

    add_heading(doc, "A.3 사용자가 직접 한 일", level=2)
    add_bullet(doc, "문제 정의: 단백질 결함 구조 분류라는 도메인 선택과 데이터 누수 비교라는 연구 질문 설정.")
    add_bullet(doc, "데이터셋 선정: Decoys ‘R’ Us의 fisa 세트 선택 (학부연구생 경험 기반).")
    add_bullet(doc, "특징 추출 코드(`src/preprocessing/extract_features.py`)와 라벨링 정책(<4Å, >8Å) 결정 및 구현.")
    add_bullet(doc, "강의 범위 한정 원칙 결정 — 학교 과제이므로 강의에서 다룬 알고리즘만 사용하기로 한 방침 자체.")
    add_bullet(doc, "실험 스크립트 실행 및 결과 검증, 보고서 최종 검수.")

    add_heading(doc, "A.4 AI가 하지 않은 일", level=2)
    add_bullet(doc, "실험 결과 수치 생성: 모든 ROC-AUC / F1 / Accuracy / Confusion Matrix 값은 본인이 직접 실행한 `python src/models/run_experiments.py`의 결과를 그대로 가져왔다. AI는 어떤 수치도 임의로 작성하지 않았다.")
    add_bullet(doc, "그림 직접 그리기: 모든 그림은 matplotlib 코드에서 자동 생성된 결과이며, AI가 별도로 이미지를 합성하지 않았다.")
    add_bullet(doc, "외부 저작물 무단 복사: 외부 코드/논문 본문을 그대로 옮긴 부분은 없다. 알고리즘 설명은 강의 자료를 1차 참고, 라이브러리 인터페이스는 공식 문서를 참고했다.")

    add_heading(doc, "A.5 검증 절차 (재현성)", level=2)
    add_paragraph(
        doc,
        "본 보고서의 모든 수치는 `random_state=42`로 고정되어 있으며, 다음 절차로 동일하게 재현된다.",
    )
    add_bullet(doc, "1. `pip install -r requirements.txt`로 환경 구성 (scikit-learn ≥ 1.4, xgboost ≥ 3.0, python-docx).")
    add_bullet(doc, "2. `python src/preprocessing/extract_features.py`로 `data/processed/features.csv` 생성 (이미 존재하면 생략 가능).")
    add_bullet(doc, "3. `python src/models/run_experiments.py`로 `results/experiment_results.json` 및 7장의 figure(`figures/`) 생성.")
    add_bullet(doc, "4. `python src/models/build_report.py`로 본 워드 보고서 재생성.")

    # ---- 부록 B. References ----
    add_heading(doc, "부록 B. References", level=1)
    add_paragraph(
        doc,
        "본문에서 인용된 자료를 인용 번호 순서로 정리한다. 강의자료([2])는 본문 곳곳에서 \"Lecture N\" 형태로 직접 인용된다.",
        size_pt=10,
    )
    refs = [
        "[1] Samudrala, R., & Levitt, M. (2000). Decoys ‘R’ Us: A database of incorrect conformations to improve protein structure prediction. Protein Science, 9(7), 1399–1401. (인용 위치: §2.1 데이터셋 출처)  http://compbio.buffalo.edu/dd/download.shtml",
        "[2] Lee, P. (2026). Machine Learning Lectures 0–10. Inha University, Multimodal AI Lab. (인용 위치: §3, §4 전반 — Lecture 2/3/6/7/10)",
        "[3] Cortes, C., & Vapnik, V. (1995). Support-vector networks. Machine Learning, 20, 273–297. (인용 위치: §3.2 SVM, §4.3 SVM-RBF 튜닝)",
        "[4] Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32. (인용 위치: §3.2 Random Forest, §5.5 MDI 기반 Feature Importance)",
        "[5] Freund, Y., & Schapire, R. E. (1997). A decision-theoretic generalization of on-line learning and an application to boosting. Journal of Computer and System Sciences, 55(1), 119–139. (인용 위치: §3.2 AdaBoost)",
        "[6] Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. In Proc. ACM SIGKDD ’16. (인용 위치: §3.2 XGBoost, §4.3 XGBoost 튜닝)",
        "[7] Pedregosa, F. et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825–2830. (인용 위치: §4 전반 — 모델 구현체 및 GroupShuffleSplit/StandardScaler/PCA 사용)",
    ]
    for r in refs:
        add_paragraph(doc, r, size_pt=10)

    return doc


def main():
    with open(RESULTS_PATH, "r", encoding="utf-8") as f:
        results = json.load(f)
    doc = build(results)
    try:
        doc.save(OUT_PATH)
        print(f"Saved: {OUT_PATH}")
    except PermissionError:
        doc.save(OUT_PATH_FALLBACK)
        print(f"(원본이 열려 있어 새 이름으로 저장) Saved: {OUT_PATH_FALLBACK}")


if __name__ == "__main__":
    main()
